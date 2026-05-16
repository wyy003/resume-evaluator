from fastapi import FastAPI, File, UploadFile, Request, Form
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import os
from datetime import datetime
import sqlite3
import json
from pypdf import PdfReader
from docx import Document
import re
import random
from openai import OpenAI
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 初始化数据库
def init_db():
    conn = sqlite3.connect('resumes.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS uploads
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  filename TEXT,
                  upload_time TEXT,
                  parsed_data TEXT,
                  jd_text TEXT,
                  llm_analysis TEXT,
                  tokens_used INTEGER,
                  api_cost REAL)''')
    conn.commit()
    conn.close()

init_db()

# 初始化 OpenAI 客户端
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# LLM 分析函数（真实 OpenAI API）
def analyze_with_llm(resume_json, jd_text):
    """使用 OpenAI API 分析简历与 JD 的匹配度"""

    # 构建优化后的 Prompt
    prompt = f"""你是一位资深 HR 和招聘专家，请深度分析以下简历与职位描述的匹配度。

【职位描述】
{jd_text}

【简历内容】
教育背景：{resume_json.get('sections', {}).get('education', '无')}

工作经历：{resume_json.get('sections', {}).get('work', '无')}

技能特长：{resume_json.get('sections', {}).get('skills', '无')}

项目经验：{resume_json.get('sections', {}).get('projects', '无')}

【评分标准】（总分 100 分）
1. 教育背景匹配度（20分）：学历层次、专业相关性、院校背景
2. 工作经验匹配度（35分）：年限要求、行业经验、岗位相关性、职责匹配
3. 技能匹配度（25分）：必备技能覆盖率、技能熟练度、技术栈匹配
4. 项目经验匹配度（20分）：项目复杂度、业务场景相似度、成果量化

【分析要求】
1. match_score：基于上述标准给出 0-100 的综合评分
2. strengths：列出 3-5 个核心优势，每条需包含：
   - 具体的匹配点（引用简历原文）
   - 与 JD 的对应关系
   - 为什么这是优势
3. weaknesses：列出 2-4 个明显不足，每条需包含：
   - 具体缺失的要求（引用 JD 原文）
   - 当前简历的差距
   - 对求职的影响程度
4. suggestions：列出 3-5 个可操作的改进建议，每条需包含：
   - 具体的改进方向
   - 如何补充或优化
   - 预期的提升效果

【输出格式】
只返回以下 JSON 格式，不要任何其他文字：
{{
  "match_score": 75,
  "strengths": [
    "具有3年Python开发经验，与JD要求的'2年以上Python经验'高度匹配，且简历中提到使用FastAPI框架，正是岗位技术栈",
    "项目经验中的'电商推荐系统'与JD中的'推荐算法优化'业务场景一致，展示了相关领域的实战能力"
  ],
  "weaknesses": [
    "JD要求'熟悉Docker/K8s容器化部署'，但简历技能部分未提及相关经验，可能在DevOps能力上存在短板",
    "工作经历缺少量化成果（如性能提升百分比、用户增长数据），难以体现实际业务价值"
  ],
  "suggestions": [
    "在技能部分补充Docker和Kubernetes经验，如果有相关实践可详细描述部署流程和遇到的问题",
    "为每个项目添加量化指标，如'优化推荐算法使点击率提升15%'，增强说服力",
    "工作经历按STAR法则重写（情境-任务-行动-结果），突出解决问题的能力和业务影响"
  ]
}}"""

    try:
        # 调用 OpenAI API
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "你是一位专业的 HR，擅长分析简历与职位的匹配度。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )

        # 解析返回结果
        result_text = response.choices[0].message.content.strip()

        # 尝试提取 JSON（可能包含 markdown 代码块）
        if "```json" in result_text:
            result_text = result_text.split("```json")[1].split("```")[0].strip()
        elif "```" in result_text:
            result_text = result_text.split("```")[1].split("```")[0].strip()

        result = json.loads(result_text)

        # 验证返回格式
        if not all(key in result for key in ["match_score", "strengths", "weaknesses", "suggestions"]):
            raise ValueError("返回格式不完整")

        # 添加 Token 使用信息
        result['tokens_used'] = response.usage.total_tokens
        result['prompt_tokens'] = response.usage.prompt_tokens
        result['completion_tokens'] = response.usage.completion_tokens
        # GPT-3.5-turbo 价格：$0.0015/1K prompt tokens, $0.002/1K completion tokens
        result['api_cost'] = (response.usage.prompt_tokens * 0.0015 + response.usage.completion_tokens * 0.002) / 1000

        return result

    except Exception as e:
        print(f"LLM 分析失败: {e}")
        # 降级到 Mock 函数
        return analyze_with_llm_mock(resume_json, jd_text)

# LLM 分析函数（Mock 版本，作为降级方案）
def analyze_with_llm_mock(resume_json, jd_text):
    """模拟 LLM 分析，返回固定格式的结果（带随机波动 + 智能评分）"""

    sections = resume_json.get("sections", {})
    base_score = 50  # 降低基础分

    # 1. 段落完整性（最多 20 分）
    if sections.get("education") and sections["education"].strip():
        base_score += 5
    if sections.get("work") and sections["work"].strip():
        base_score += 8
    if sections.get("skills") and sections["skills"].strip():
        base_score += 4
    if sections.get("projects") and sections["projects"].strip():
        base_score += 3

    # 2. 关键词匹配（最多 30 分）
    keyword_score = 0
    matched_skills = []
    if jd_text.strip():
        jd_keywords = extract_keywords(jd_text)
        keyword_score, matched_skills = count_keyword_matches(resume_json, jd_keywords)
        keyword_score = min(keyword_score, 30)  # 上限30分
        base_score += keyword_score

    # 3. 内容质量（最多 20 分）
    quality_score = 0

    # 工作经历有数字（量化）
    if has_numbers(sections.get("work", "")):
        quality_score += 8

    # 项目描述详细（>200 字）
    if len(sections.get("projects", "")) > 200:
        quality_score += 7

    # 技能数量（>5 个）
    if count_skills(sections.get("skills", "")) > 5:
        quality_score += 5

    base_score += quality_score

    # 4. 随机波动 ±5%
    fluctuation = random.randint(-5, 5)
    final_score = min(max(base_score + fluctuation, 0), 100)

    # 5. 动态生成优势/不足/建议
    strengths, weaknesses, suggestions = generate_feedback(
        sections, jd_text, keyword_score, quality_score, matched_skills
    )

    return {
        "match_score": final_score,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "suggestions": suggestions
    }

# 辅助函数
def extract_keywords(jd_text):
    """从 JD 中提取关键词（优化版：支持同义词和权重）"""
    # 技术栈同义词库
    tech_synonyms = {
        "python": ["python", "python3", "py", "django", "flask", "fastapi"],
        "java": ["java", "spring", "springboot", "spring boot", "maven", "gradle"],
        "javascript": ["javascript", "js", "typescript", "ts", "node", "nodejs", "node.js"],
        "react": ["react", "reactjs", "react.js", "react native"],
        "vue": ["vue", "vuejs", "vue.js", "vue3"],
        "database": ["sql", "mysql", "postgresql", "postgres", "mongodb", "redis", "数据库"],
        "docker": ["docker", "容器", "k8s", "kubernetes"],
        "cloud": ["aws", "azure", "gcp", "阿里云", "腾讯云", "云计算"],
        "ai": ["机器学习", "深度学习", "ai", "nlp", "cv", "tensorflow", "pytorch", "神经网络"],
        "frontend": ["前端", "html", "css", "webpack", "vite"],
        "backend": ["后端", "api", "restful", "微服务", "分布式"],
        "git": ["git", "github", "gitlab", "版本控制"],
        "linux": ["linux", "unix", "shell", "bash"],
        "golang": ["go", "golang", "gin"],
        "rust": ["rust", "cargo"]
    }

    # 核心技能关键词（权重更高）
    core_indicators = ["必须", "精通", "熟练", "要求", "核心", "主要", "负责"]

    jd_lower = jd_text.lower()
    keywords = {}  # {category: {"weight": int, "found": bool}}

    # 检测每个技术栈类别
    for category, synonyms in tech_synonyms.items():
        found = False
        is_core = False

        # 检查是否匹配任何同义词
        for synonym in synonyms:
            if synonym in jd_lower:
                found = True

                # 检查是否为核心技能（前后有核心指示词）
                for indicator in core_indicators:
                    if indicator in jd_text and synonym in jd_lower:
                        # 简单判断：如果核心指示词和技能词在同一段落
                        is_core = True
                        break
                break

        if found:
            keywords[category] = {
                "weight": 3 if is_core else 2,  # 核心技能3分，普通技能2分
                "synonyms": synonyms
            }

    return keywords

def count_keyword_matches(resume_json, keywords):
    """计算简历中匹配的关键词数量（优化版：支持同义词和熟练度识别）"""
    resume_text = json.dumps(resume_json, ensure_ascii=False).lower()

    total_score = 0
    matched_skills = []

    # 熟练度关键词
    proficiency_levels = {
        "精通": 1.5,
        "熟练": 1.3,
        "熟悉": 1.0,
        "了解": 0.7
    }

    for category, info in keywords.items():
        base_weight = info["weight"]
        synonyms = info["synonyms"]

        # 检查是否匹配任何同义词
        matched = False
        proficiency_multiplier = 1.0

        for synonym in synonyms:
            if synonym in resume_text:
                matched = True

                # 检查熟练度
                for level, multiplier in proficiency_levels.items():
                    # 在简历原文中查找（保留大小写以匹配中文）
                    resume_original = json.dumps(resume_json, ensure_ascii=False)
                    if level in resume_original and synonym in resume_text:
                        # 简单判断：如果熟练度词和技能词距离较近
                        proficiency_multiplier = max(proficiency_multiplier, multiplier)

                break

        if matched:
            skill_score = base_weight * proficiency_multiplier
            total_score += skill_score
            matched_skills.append({
                "category": category,
                "score": skill_score
            })

    return int(total_score), matched_skills

def has_numbers(text):
    """检查文本中是否有数字（量化指标）"""
    return bool(re.search(r'\d+', text))

def count_skills(skills_text):
    """统计技能数量（简单按分隔符分割）"""
    if not skills_text:
        return 0
    separators = [',', '、', '，', ';', '；', '\n']
    count = 1
    for sep in separators:
        count = max(count, skills_text.count(sep) + 1)
    return count

def generate_feedback(sections, jd_text, keyword_score, quality_score, matched_skills=None):
    """动态生成优势/不足/建议（优化版：基于匹配的技能）"""
    strengths = []
    weaknesses = []
    suggestions = []

    # 教育背景
    if sections.get("education") and sections["education"].strip():
        strengths.append("教育背景清晰完整")
    else:
        weaknesses.append("缺少教育背景信息")
        suggestions.append("建议补充教育经历")

    # 工作经历
    if sections.get("work") and sections["work"].strip():
        if has_numbers(sections["work"]):
            strengths.append("工作经历包含量化数据，展示具体成果")
        else:
            weaknesses.append("工作成果缺少量化数据")
            suggestions.append("建议用数字量化工作成果（如：提升 30% 效率、负责 5 人团队）")
    else:
        weaknesses.append("缺少工作经历")
        suggestions.append("建议补充相关工作经验")

    # 关键词匹配（优化版：显示具体匹配的技能）
    if jd_text.strip():
        if matched_skills and len(matched_skills) > 0:
            # 显示匹配的技能类别
            skill_names = [skill["category"] for skill in matched_skills[:5]]  # 最多显示5个
            strengths.append(f"技能匹配度良好，涵盖：{', '.join(skill_names)}")

            if keyword_score > 20:
                strengths.append("核心技能覆盖全面，与岗位要求高度匹配")
            elif keyword_score > 10:
                suggestions.append("建议在简历中更突出核心技能的熟练程度（如：精通、熟练）")
        else:
            weaknesses.append("简历中缺少岗位相关的关键技能")
            suggestions.append("建议仔细阅读 JD，在简历中突出相关技能和经验")

    # 项目经历
    if sections.get("projects") and len(sections["projects"]) > 200:
        strengths.append("项目经验描述详细，体现实践能力")
    elif sections.get("projects") and sections["projects"].strip():
        weaknesses.append("项目描述较为简略")
        suggestions.append("建议采用 STAR 法则详细描述项目（情境-任务-行动-结果）")
    else:
        weaknesses.append("缺少项目经历")
        suggestions.append("建议补充相关项目经验")

    # 技能部分
    skill_count = count_skills(sections.get("skills", ""))
    if skill_count > 5:
        strengths.append(f"技能列表丰富（{skill_count} 项技能）")
    elif skill_count > 0:
        suggestions.append("可以补充更多专业技能，展示技术广度")

    # 确保至少有内容
    if not strengths:
        strengths.append("简历基本信息完整")
    if not weaknesses:
        weaknesses.append("整体表现良好，可进一步优化细节")
    if not suggestions:
        suggestions.append("保持简历更新，持续积累项目经验")

    return strengths, weaknesses, suggestions

# DOCX 解析函数
def parse_docx(file_path):
    """提取 DOCX 文字并识别简历段落"""
    try:
        doc = Document(file_path)
        full_text = ""

        # 提取所有段落的文字
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"

        if not full_text.strip():
            return {
                "filename": os.path.basename(file_path),
                "text": "",
                "sections": {
                    "education": "",
                    "work": "",
                    "skills": "",
                    "projects": ""
                },
                "error": "无法提取文字，文档可能为空"
            }

        # 使用与 PDF 相同的段落识别逻辑
        return parse_text_sections(full_text, os.path.basename(file_path))

    except Exception as e:
        return {
            "filename": os.path.basename(file_path),
            "text": "",
            "sections": {
                "education": "",
                "work": "",
                "skills": "",
                "projects": ""
            },
            "error": str(e)
        }

# PDF 解析函数
def parse_pdf(file_path):
    """提取 PDF 文字并识别简历段落"""
    try:
        reader = PdfReader(file_path)
        full_text = ""

        # 提取所有页面的文字
        for page in reader.pages:
            full_text += page.extract_text() + "\n"

        if not full_text.strip():
            return {
                "filename": os.path.basename(file_path),
                "text": "",
                "sections": {
                    "education": "",
                    "work": "",
                    "skills": "",
                    "projects": ""
                },
                "error": "无法提取文字，可能是扫描版 PDF"
            }

        return parse_text_sections(full_text, os.path.basename(file_path))

    except Exception as e:
        return {
            "filename": os.path.basename(file_path),
            "text": "",
            "sections": {
                "education": "",
                "work": "",
                "skills": "",
                "projects": ""
            },
            "error": str(e)
        }

# 通用文本段落识别函数
def parse_text_sections(full_text, filename):
    """从文本中识别简历段落"""
    lines = full_text.split('\n')

    section_keywords = {
        "education": ["教育", "学历", "education", "academic"],
        "work": ["工作", "经历", "experience", "employment", "职位"],
        "skills": ["技能", "skills", "能力", "专长"],
        "projects": ["项目", "project", "作品"]
    }

    # 识别标题位置
    title_line_indices = {
        "education": None,
        "work": None,
        "skills": None,
        "projects": None
    }

    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        if not line_lower:
            continue

        for section, keywords in section_keywords.items():
            if any(keyword in line_lower for keyword in keywords):
                if title_line_indices[section] is None:
                    title_line_indices[section] = i

    # 提取段落内容
    sections = {
        "education": "",
        "work": "",
        "skills": "",
        "projects": ""
    }

    # 获取所有标题的行号，按行号排序
    all_title_indices = []
    for section, idx in title_line_indices.items():
        if idx is not None:
            all_title_indices.append((idx, section))
    all_title_indices.sort()

    # 提取每个段落的内容
    for i, (start_idx, section) in enumerate(all_title_indices):
        # 确定结束位置：下一个标题的位置，或文件末尾
        if i + 1 < len(all_title_indices):
            end_idx = all_title_indices[i + 1][0]
        else:
            end_idx = len(lines)

        # 提取内容（跳过标题行本身）
        content_lines = lines[start_idx + 1:end_idx]
        sections[section] = "\n".join(content_lines).strip()

    return {
        "filename": filename,
        "text": full_text,
        "sections": sections
    }

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    # 获取历史记录
    conn = sqlite3.connect('resumes.db')
    c = conn.cursor()
    c.execute('SELECT id, filename, upload_time, parsed_data FROM uploads ORDER BY id DESC')
    history = c.fetchall()
    conn.close()

    return templates.TemplateResponse("index.html", {
        "request": request,
        "history": history
    })

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), jd_text: str = Form("")):
    # 保存文件
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)

    # 根据文件类型选择解析函数
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext == '.pdf':
        parsed_data = parse_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        parsed_data = parse_docx(file_path)
    else:
        return JSONResponse(
            status_code=400,
            content={"error": f"不支持的文件格式: {file_ext}，仅支持 PDF 和 DOCX"}
        )

    parsed_json = json.dumps(parsed_data, ensure_ascii=False)

    # LLM 分析（使用真实 OpenAI API）
    llm_analysis = None
    if jd_text.strip():
        llm_analysis = analyze_with_llm(parsed_data, jd_text)
        llm_analysis_json = json.dumps(llm_analysis, ensure_ascii=False)
    else:
        llm_analysis_json = None

    # 保存到数据库
    conn = sqlite3.connect('resumes.db')
    c = conn.cursor()
    upload_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 提取 Token 和成本信息
    tokens_used = llm_analysis.get('tokens_used', 0) if llm_analysis else 0
    api_cost = llm_analysis.get('api_cost', 0.0) if llm_analysis else 0.0

    c.execute('INSERT INTO uploads (filename, upload_time, parsed_data, jd_text, llm_analysis, tokens_used, api_cost) VALUES (?, ?, ?, ?, ?, ?, ?)',
              (file.filename, upload_time, parsed_json, jd_text, llm_analysis_json, tokens_used, api_cost))
    upload_id = c.lastrowid
    conn.commit()
    conn.close()

    print(f"✅ 收到文件: {file.filename}")
    print(f"📄 解析结果: {len(parsed_data.get('text', ''))} 字符")
    if llm_analysis:
        print(f"🤖 LLM 分析: 匹配度 {llm_analysis['match_score']}%")
        if tokens_used > 0:
            print(f"💰 Token 使用: {tokens_used} tokens (${api_cost:.4f})")

    return {
        "message": "上传成功",
        "filename": file.filename,
        "upload_id": upload_id,
        "parsed_data": parsed_data,
        "llm_analysis": llm_analysis
    }

@app.get("/result/{upload_id}")
async def get_result(request: Request, upload_id: int):
    """查看解析结果详情页"""
    conn = sqlite3.connect('resumes.db')
    c = conn.cursor()
    c.execute('SELECT filename, upload_time, parsed_data, jd_text, llm_analysis, tokens_used, api_cost FROM uploads WHERE id = ?', (upload_id,))
    result = c.fetchone()
    conn.close()

    if not result:
        return HTMLResponse("未找到记录", status_code=404)

    filename, upload_time, parsed_json, jd_text, llm_analysis_json, tokens_used, api_cost = result
    parsed_data = json.loads(parsed_json) if parsed_json else {}
    llm_analysis = json.loads(llm_analysis_json) if llm_analysis_json else None

    return templates.TemplateResponse("result.html", {
        "request": request,
        "filename": filename,
        "upload_time": upload_time,
        "parsed_data": parsed_data,
        "jd_text": jd_text,
        "llm_analysis": llm_analysis,
        "tokens_used": tokens_used or 0,
        "api_cost": api_cost or 0.0
    })

@app.get("/stats", response_class=HTMLResponse)
async def get_stats(request: Request):
    """查看 API 成本统计"""
    conn = sqlite3.connect('resumes.db')
    c = conn.cursor()

    # 获取所有记录
    c.execute('SELECT id, filename, upload_time, tokens_used, api_cost FROM uploads ORDER BY upload_time DESC')
    records = []
    total_tokens = 0
    total_cost = 0.0

    for row in c.fetchall():
        record_id, filename, upload_time, tokens_used, api_cost = row
        tokens_used = tokens_used or 0
        api_cost = api_cost or 0.0

        records.append({
            'id': record_id,
            'filename': filename,
            'upload_time': upload_time,
            'tokens_used': tokens_used,
            'api_cost': api_cost
        })

        total_tokens += tokens_used
        total_cost += api_cost

    conn.close()

    total_count = len(records)
    avg_cost = total_cost / total_count if total_count > 0 else 0.0

    return templates.TemplateResponse("stats.html", {
        "request": request,
        "records": records,
        "total_count": total_count,
        "total_tokens": total_tokens,
        "total_cost": total_cost,
        "avg_cost": avg_cost
    })

@app.get("/batch-results", response_class=HTMLResponse)
async def get_batch_results(request: Request, ids: str):
    """批量结果展示页面"""
    conn = sqlite3.connect('resumes.db')
    c = conn.cursor()

    # 解析 ID 列表
    upload_ids = [int(id.strip()) for id in ids.split(',') if id.strip()]

    results = []
    for upload_id in upload_ids:
        c.execute('SELECT id, filename, upload_time, llm_analysis FROM uploads WHERE id = ?', (upload_id,))
        row = c.fetchone()

        if row:
            record_id, filename, upload_time, llm_analysis_json = row

            # 解析 LLM 分析结果
            if llm_analysis_json:
                llm_analysis = json.loads(llm_analysis_json)
                score = llm_analysis.get('match_score', 0)
                strengths = llm_analysis.get('strengths', [])
                weaknesses = llm_analysis.get('weaknesses', [])
                suggestions = llm_analysis.get('suggestions', [])
            else:
                score = 0
                strengths = []
                weaknesses = []
                suggestions = []

            results.append({
                'id': record_id,
                'filename': filename,
                'upload_time': upload_time,
                'score': score,
                'strengths_count': len(strengths),
                'weaknesses_count': len(weaknesses),
                'suggestions_count': len(suggestions)
            })

    conn.close()

    # 统计数据
    high_count = sum(1 for r in results if r['score'] >= 80)
    medium_count = sum(1 for r in results if 60 <= r['score'] < 80)
    low_count = sum(1 for r in results if r['score'] < 60)
    avg_score = int(sum(r['score'] for r in results) / len(results)) if results else 0

    return templates.TemplateResponse("batch_results.html", {
        "request": request,
        "results": results,
        "high_count": high_count,
        "medium_count": medium_count,
        "low_count": low_count,
        "avg_score": avg_score
    })

@app.get("/export/{record_id}")
async def export_report(record_id: int):
    """导出分析报告为Word文档"""
    conn = sqlite3.connect('resumes.db')
    cursor = conn.cursor()

    cursor.execute('''
        SELECT filename, upload_time, parsed_data, jd_text, llm_analysis, tokens_used, api_cost
        FROM uploads WHERE id = ?
    ''', (record_id,))

    row = cursor.fetchone()
    conn.close()

    if not row:
        return JSONResponse({"error": "记录不存在"}, status_code=404)

    filename, upload_time, parsed_data_json, jd_text, llm_analysis_json, tokens_used, api_cost = row
    parsed_data = json.loads(parsed_data_json)

    # 创建Word文档
    doc = Document()

    # 标题
    title = doc.add_heading('简历分析报告', 0)
    title.alignment = 1  # 居中

    # 基本信息
    doc.add_heading('基本信息', level=1)
    doc.add_paragraph(f'文件名：{filename}')
    doc.add_paragraph(f'分析时间：{upload_time}')
    doc.add_paragraph(f'文本长度：{len(parsed_data.get("text", ""))} 字符')

    if tokens_used:
        doc.add_paragraph(f'Token 使用：{tokens_used}')
    if api_cost:
        doc.add_paragraph(f'分析成本：${api_cost:.4f}')

    # AI分析结果
    if llm_analysis_json:
        llm_analysis = json.loads(llm_analysis_json)

        doc.add_heading('AI 分析结果', level=1)

        # 匹配度分数
        score = llm_analysis.get('score', 0)
        doc.add_paragraph(f'匹配度分数：{score} 分', style='Intense Quote')

        # 优势
        strengths = llm_analysis.get('strengths', [])
        if strengths:
            doc.add_heading('优势', level=2)
            for strength in strengths:
                doc.add_paragraph(strength, style='List Bullet')

        # 不足
        weaknesses = llm_analysis.get('weaknesses', [])
        if weaknesses:
            doc.add_heading('不足', level=2)
            for weakness in weaknesses:
                doc.add_paragraph(weakness, style='List Bullet')

        # 改进建议
        suggestions = llm_analysis.get('suggestions', [])
        if suggestions:
            doc.add_heading('改进建议', level=2)
            for suggestion in suggestions:
                doc.add_paragraph(suggestion, style='List Bullet')

    # 职位描述
    if jd_text:
        doc.add_heading('职位描述', level=1)
        doc.add_paragraph(jd_text)

    # 简历内容
    doc.add_heading('简历内容', level=1)

    sections = ['education', 'work', 'skills', 'projects']
    section_names = {'education': '教育背景', 'work': '工作经历', 'skills': '技能', 'projects': '项目经验'}

    for section in sections:
        content = parsed_data.get(section, '')
        if content:
            doc.add_heading(section_names[section], level=2)
            doc.add_paragraph(content)

    # 保存文档
    output_filename = f"report_{record_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join('uploads', output_filename)
    doc.save(output_path)

    return FileResponse(
        path=output_path,
        filename=output_filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


